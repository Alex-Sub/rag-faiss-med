from __future__ import annotations

import os
import json
from typing import Iterable, Tuple, Optional

import fitz  # PyMuPDF
from bs4 import BeautifulSoup
from docx import Document


from utils import chunk_text, clean_text

ROOT = os.path.dirname(os.path.dirname(__file__))
DOCS_DIR = os.path.join(ROOT, "documents")
OUT_PATH = os.path.join(ROOT, "processed", "chunks.jsonl")

# какие расширения индексируем
ALLOWED_EXT = {".pdf", ".html", ".htm", ".txt", ".docx"}
# что игнорируем
SKIP_EXT = {".zip", ".7z", ".rar", ".exe", ".dll"}

def iter_files(root_dir: str) -> Iterable[str]:
    for root, _, files in os.walk(root_dir):
        for fn in files:
            path = os.path.join(root, fn)
            ext = os.path.splitext(fn)[1].lower()
            if ext in SKIP_EXT:
                continue
            if ext in ALLOWED_EXT:
                yield path

def iter_pdf_pages(path: str) -> Iterable[Tuple[int, str]]:
    """Yield (page_number_1based, text)."""
    doc = fitz.open(path)
    for i in range(doc.page_count):
        page = doc.load_page(i)
        txt = page.get_text("text") or ""
        yield (i + 1), txt

def html_to_text(path: str) -> str:
    with open(path, "rb") as f:
        soup = BeautifulSoup(f.read(), "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(separator="\n")

def read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def read_docx(path: str) -> str:
    doc = Document(path)
    # абзацы
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # таблицы (часто важны в медицине)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n".join(paras)

def write_rec(out_f, rec: dict) -> None:
    out_f.write((json.dumps(rec, ensure_ascii=False) + "\n").encode("utf-8"))
def get_short_path(path: str) -> str:
    """Return Windows 8.3 short path to avoid Unicode issues in some native libs."""
    GetShortPathNameW = ctypes.windll.kernel32.GetShortPathNameW
    GetShortPathNameW.argtypes = [wintypes.LPCWSTR, wintypes.LPWSTR, wintypes.DWORD]
    GetShortPathNameW.restype = wintypes.DWORD

    buf = ctypes.create_unicode_buffer(4096)
    res = GetShortPathNameW(path, buf, 4096)
    if res == 0:
        # Если 8.3 имена отключены или не получилось — вернём исходный путь
        return path
    return buf.value

def main():
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)

    files = list(iter_files(DOCS_DIR))
    print(f"Documents root: {DOCS_DIR}")
    print(f"Found eligible files: {len(files)}")

    by_ext = {}
    for p in files:
        ext = os.path.splitext(p)[1].lower()
        by_ext[ext] = by_ext.get(ext, 0) + 1
    print("By type:", ", ".join([f"{k}={v}" for k, v in sorted(by_ext.items())]) or "none")

    count = 0
    skipped_empty_pages = 0

    with open(OUT_PATH, "wb") as out:
        for path in files:
            fname = os.path.basename(path)
            ext = os.path.splitext(path)[1].lower()

            # PDF: чанки по страницам (для цитат "страница")
            if ext == ".pdf":
                for page_num, txt in iter_pdf_pages(path):
                    txt = clean_text(txt)
                    # если текст пустой — это может быть скан (OCR будет позже)
                    if len(txt) < 20:
                        skipped_empty_pages += 1
                        continue

                    chunks = chunk_text(txt)
                    for j, ch in enumerate(chunks, start=1):
                        ch = clean_text(ch)
                        if not ch:
                            continue
                        rec = {
                            "id": f"{fname}::p{page_num}::c{j}",
                            "text": ch,
                            "source_file": fname,
                            "page": page_num,
                            "chunk_in_page": j,
                            "type": "pdf",
                        }
                        write_rec(out, rec)
                        count += 1

            # HTML / HTM
            elif ext in (".html", ".htm"):
                txt = clean_text(html_to_text(path))
                if len(txt) < 20:
                    continue
                chunks = chunk_text(txt)
                for j, ch in enumerate(chunks, start=1):
                    ch = clean_text(ch)
                    if not ch:
                        continue
                    rec = {
                        "id": f"{fname}::c{j}",
                        "text": ch,
                        "source_file": fname,
                        "page": None,
                        "chunk_in_page": j,
                        "type": "html",
                    }
                    write_rec(out, rec)
                    count += 1

            # TXT
            elif ext == ".txt":
                txt = clean_text(read_txt(path))
                if len(txt) < 20:
                    continue
                chunks = chunk_text(txt)
                for j, ch in enumerate(chunks, start=1):
                    ch = clean_text(ch)
                    if not ch:
                        continue
                    rec = {
                        "id": f"{fname}::c{j}",
                        "text": ch,
                        "source_file": fname,
                        "page": None,
                        "chunk_in_page": j,
                        "type": "txt",
                    }
                    write_rec(out, rec)
                    count += 1

            # DOCX
            elif ext == ".docx":
                txt = clean_text(read_docx(path))
                if len(txt) < 20:
                    continue
                chunks = chunk_text(txt)
                for j, ch in enumerate(chunks, start=1):
                    ch = clean_text(ch)
                    if not ch:
                        continue
                    rec = {
                        "id": f"{fname}::c{j}",
                        "text": ch,
                        "source_file": fname,
                        "page": None,
                        "chunk_in_page": j,
                        "type": "docx",
                    }
                    write_rec(out, rec)
                    count += 1

    print(f"OK. chunks written: {count}")
    print(f"Output: {OUT_PATH}")
    if skipped_empty_pages:
        print(f"PDF pages skipped as empty/scan-like: {skipped_empty_pages} (OCR will handle later)")

if __name__ == "__main__":
    main()
